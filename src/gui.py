from PyQt6.QtCore import QSize, Qt
from PyQt6.QtWidgets import (QApplication, QMainWindow, QDialog, QPushButton, QLabel,
                             QLineEdit, QComboBox, QVBoxLayout, QWidget, QHBoxLayout, QDialogButtonBox)
from classes.Burger import Burger
from classes.Pizza import Pizza
from classes.vok import Vok
from docx import Document
from openpyxl import Workbook

class AlertBox(QDialog):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Ошибка!")

        q_btn = QDialogButtonBox.StandardButton.Ok
        self.button_box = QDialogButtonBox(q_btn)
        self.button_box.accepted.connect(self.accept)

        layout = QVBoxLayout()
        message = QLabel("Введено не число!")
        layout.addWidget(message)
        layout.addWidget(self.button_box)
        self.setLayout(layout)

class MainWindow(QMainWindow):
    def __init__(self) -> None:
        super(MainWindow, self).__init__()
        self.result_label = None
        self.cost_textbox = None
        self.width_textbox = None
        self.combobox = None
        self.setWindowTitle("MyApp")
        self.setFixedSize(QSize(800, 400))
        layout = self.initialize_out_layout()
        widget = QWidget()
        widget.setLayout(layout)
        self.setCentralWidget(widget)

    def initialize_out_layout(self) -> QVBoxLayout:
        layout = QVBoxLayout()
        layout2 = QVBoxLayout()
        horizontal_layout = QHBoxLayout()
        horizontal_layout2 = QHBoxLayout()
        buttons_layout = QHBoxLayout()
        layout.setSpacing(30)
        layout2.setSpacing(30)

        label = QLabel("Энергетическая ценность и стоимость:")
        font = label.font()
        font.setPointSize(20)
        label.setFont(font)
        # label.setStyleSheet("""padding-bottom: 10px;""")
        layout.addWidget(label, alignment=Qt.AlignmentFlag.AlignHCenter)

        self.combobox = QComboBox()
        self.combobox.setFixedSize(500, 50)
        self.combobox.addItems(["Вок", "Пицка", "Бургер"])
        self.combobox.setCurrentIndex(0)
        layout.addWidget(self.combobox, alignment=Qt.AlignmentFlag.AlignHCenter)

        self.initialize_in_layout(horizontal_layout, horizontal_layout2, layout2)

        layout.addLayout(layout2)

        self.result_label = QLabel()
        self.result_label.setFixedSize(800, 40)
        font = self.result_label.font()
        font.setPointSize(15)
        self.result_label.setFont(font)
        self.result_label.setAlignment(Qt.AlignmentFlag.AlignHCenter)
        layout.addWidget(self.result_label, alignment=Qt.AlignmentFlag.AlignHCenter)


        button = QPushButton("Рассчёт")
        button.setFixedSize(100, 40)
        button.clicked.connect(self.button_click)
        buttons_layout.addWidget(button)

        button_to_doc = QPushButton(".docx")
        button_to_doc.setFixedSize(100, 40)
        button_to_doc.clicked.connect(self.to_docx)
        buttons_layout.addWidget(button_to_doc)

        button_to_xls = QPushButton(".xlsx")
        button_to_xls.setFixedSize(100, 40)
        button_to_xls.clicked.connect(self.to_xls)
        buttons_layout.addWidget(button_to_xls)

        layout.addLayout(buttons_layout)

        return layout


    def initialize_in_layout(self, horizontal_layout : QHBoxLayout, horizontal_layout2 : QHBoxLayout,
                             layout2 : QVBoxLayout) -> None:
        self.initialize_horizontal_layout(horizontal_layout)

        layout2.addLayout(horizontal_layout)

        self.initialize_horizontal_layout2(horizontal_layout2)

        layout2.addLayout(horizontal_layout2)

    def initialize_horizontal_layout(self, horizontal_layout : QHBoxLayout) -> None:
        self.width_textbox = QLineEdit()
        self.width_textbox.setFixedSize(375, 40)
        self.width_textbox.setPlaceholderText("Введите энерг.цен. ингредиентов:")
        horizontal_layout.addWidget(self.width_textbox)

    def initialize_horizontal_layout2(self, horizontal_layout2 : QHBoxLayout) -> None:
        self.cost_textbox = QLineEdit()
        self.cost_textbox.setFixedSize(375, 40)
        self.cost_textbox.setPlaceholderText("Введите стоимости ингредиентов:")
        horizontal_layout2.addWidget(self.cost_textbox)

    def button_click(self) -> None:
        try:
           en_cost = [float(x.strip()) for x in self.width_textbox.text().split(',')]
           mn_cost = [float(x.strip()) for x in self.cost_textbox.text().split(',')]
           print (en_cost)
           print (mn_cost)
           if en_cost==[] or mn_cost==[]:
               dlg = AlertBox()
               dlg.exec()
               return
        except ValueError:
            dlg = AlertBox()
            dlg.exec()
            return

        index = self.combobox.currentText()
        obj = None
        match index:
            case "Вок":
                obj = Vok()
            case "Пицка":
                obj = Pizza()
            case "Бургер":
                obj = Burger()
        obj.calculate_energy_cost(en_cost)
        obj.calculate_money_cost(mn_cost)
        self.result_label.setText(str(obj))

    def to_docx(self) -> None:
        if self.result_label.text() == "":
            return;
        document = Document()
        document.add_heading('Энергетическая ценность и стоимость:', 0)
        document.add_paragraph(self.result_label.text())
        document.save("word.docx")

    def to_xls(self) -> None:
        if self.result_label.text() == "":
            return

        wb = Workbook()
        ws = wb.active
        ws['A1'] = 'Энергетическая ценность и стоимость:'
        ws['A2'] = self.result_label.text()
        wb.save("excel.xlsx")



app = QApplication([])

window = MainWindow()
window.show()

app.exec()